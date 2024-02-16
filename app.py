import os
import json5 as json
from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm 

# Mapping of templates for different report sections
template_map = {
    'Full Standard Due Diligence': {
        'cover_page': 'templates/Cover Page - Standard.docx',
        'cover_b': 'templates/Cover Page B.docx',
        'matrix': 'templates/Matrix.docx',
        'section': 'templates/Section Heading.docx',
        'Real Estate Ownership': 'templates/components/Real Estate Ownership.docx',
        'Bankruptcies & Adversary Proceedings': 'templates/components/Bankruptcy.docx',
        'Lien': 'templates/components/Lien.docx',
        'County, State Criminal Case': 'templates/components/County, State Criminal Case.docx'
    }
}

icons = {
    "NO_RECORDS": "templates/graphics/checkmark.png",
    "RECORDS": "templates/graphics/red_triangle.png",
    "NOT_APPLICABLE": "-",
    "PENDING_RECORDS": "templates/graphics/pending.png",
    "DOC_RETRIEVAL": "templates/graphics/mailbox.png",
}

class DocGenerator:
    def __init__(self, report_json_file, report_file):
        """Initialize the document generator with JSON data and output file."""
        self.report_json_data = self.read_json(report_json_file)
        self.report_file = report_file
        self.doc = Document()
        self.load_metadata()
        
    def set_column_widths(self, table, widths):
        """Set table column widths."""
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                # Set cell width
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(widths[idx].twips))  # Set width
                tcW.set(qn('w:type'), 'dxa')  # Set type as dxa (twentieths of a point)
                cell._element.get_or_add_tcPr().append(tcW)  # Append width element to cell properties

    def insert_bookmark_start(self, paragraph, bookmark_id, bookmark_name):
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), bookmark_name)
        paragraph._p.addnext(bookmark_start)
        return bookmark_start

    def insert_bookmark_end(self, paragraph, bookmark_id):
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))
        paragraph._p.addnext(bookmark_end)
        return bookmark_end

    def generate_table_with_template(self, composer, headings, contents, sections):
        """Generate a table using docxtpl with the provided headings and contents, including images."""
        # Load your docx template that has the table structure predefined
        doc_tpl = DocxTemplate('templates/Matrix.docx')
        # Prepare contents with InlineImage objects
        updated_contents = []
        for item in contents:
            updated_columns = []
            for i, icon_path in enumerate(item['columns']):
                # Check if the value is a path to an image
                if icon_path in icons and icons[icon_path] not in ["-", "NOT_APPLICABLE"]:
                    # Convert the path to an InlineImage object
                    image = InlineImage(doc_tpl, icons[icon_path], width=Cm(0.5))  # Adjust width as needed
                    updated_columns.append(image)
                else:
                    # If not an image path, use the value directly (e.g., for placeholders like "-")
                    updated_columns.append("-")
            updated_contents.append({'label': item['label'], 'columns': updated_columns})

        # Render the document with the headings and updated contents including InlineImage objects
        context = {
            'headings': headings,
            'contents': updated_contents
        }
        doc_tpl.render(context)
        
        # Add the rendered document to the Composer
        rendered_doc = doc_tpl.docx
        
        # Assuming the table is the first one in the document
        table = rendered_doc.tables[0]

        # ToDo
        # bookmark_id = 0
        # for row in table.rows[1:]:  # Skip header row
        #     for cell in row.cells:
        #         paragraph = cell.paragraphs[0]
        #         print (type(paragraph.runs[0].element))
        #         # Create unique bookmark names based on a counter
        #         bookmark_name = f"Real Estate Holdings"
        #         self.insert_bookmark_start(paragraph, bookmark_id, bookmark_name)
        #         self.insert_bookmark_end(paragraph, bookmark_id)
        #         bookmark_id += 1


        max_head_length = 2.75 
        max_head_length = [min(max_head_length, len(heading)/4.0) for heading in headings[1:]]
        column_widths = [Cm(min(max_head_length)) for _ in range(len(headings))]
        column_widths[0] = Cm(4.5)
        # Set column widths
        self.set_column_widths(table, column_widths)
        
        composer.append(rendered_doc)
        
    def handle_special_chars(self, data):
        if isinstance(data, str):
            return data.replace('&', '&amp;')
        elif isinstance(data, dict):
            return {k: self.handle_special_chars(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self.handle_special_chars(item) for item in data]
        else:
            return data
        
    def read_json(self, filename):
        """Read and return JSON data from a file."""
        json_data = {}
        if os.path.isfile(filename):
            try:
                with open(filename, "r") as file:
                    json_data = json.load(file)
            except OSError as e:
                raise
        if json_data:
            json_data = self.handle_special_chars(json_data)
        return json_data

    def load_metadata(self):
        """Load metadata from JSON data into instance variables."""
        self.template_name = self.report_json_data.get("template_name")
        self.record_date = self.report_json_data.get("report_date")
        self.subject_name = self.report_json_data.get("subject_name")

    def generate(self):
        """Generate the final document based on the JSON data and templates."""
        template_files_map = template_map.get(self.template_name)

        # Start with the cover page as the base for the final document
        cover_page_doc = DocxTemplate(template_files_map.get('cover_page'))
        cover_page_doc.render(self.report_json_data)
        final_doc = cover_page_doc.docx  # Use the internal docx Document of cover_page_doc as the starting point

        composer = Composer(final_doc)  # Initialize Composer with the cover page document
        
        sections_data = self.report_json_data['sections']
        headings = [section_detail['table_header'] for section_name, section_detail in sections_data.items()]
        headings = [''] + headings
        jurisdictions = sections_data['Real Estate Holdings']['jurisdictions']
        
        contents = []
        for r in jurisdictions:
            contents.append({'label': r.replace(', ', ',\n'), 'columns': [sections_data[col]['jurisdictions'][r] for col in list(sections_data.keys())]})
        self.generate_table_with_template(composer, headings, contents, list(self.report_json_data.get('sections', {}).keys()))
        
        # Add cover - B
        cover_page_b_tpl = DocxTemplate(template_files_map.get('cover_b'))
        cover_page_b_tpl.render(self.report_json_data)
        
        cover_page_b_doc = cover_page_b_tpl.docx
        composer.append(cover_page_b_doc)  # Use the internal docx Document of cover_page_doc as the starting point


        # Iterate through sections and components to build the document
        for section, section_data in self.report_json_data.get('sections', {}).items():
            if section_data.get('components'):
                section_template = template_files_map.get('section')
                if section_template:
                    section_doc_tpl = DocxTemplate(section_template)
                    section_title = {'section_title': section}
                    section_doc_tpl.render(section_title)
                    section_doc = section_doc_tpl.docx  # Convert to python-docx Document object for Composer
                    composer.append(section_doc)  # Append section document

                for component in section_data.get('components', {}):
                    records = section_data["components"][component]
                    
                    component_template = template_files_map.get(component)
                    if component_template:
                        component_doc_tpl = DocxTemplate(component_template)
                        component_doc_tpl.render(records)
                        component_doc = component_doc_tpl.docx  # Convert to python-docx Document object for Composer
                        composer.append(component_doc)  # Append component document

        # Save the final merged document
        composer.doc.save(self.report_file)

# Example usage
json_file = "test_report.json"
out_file = "test_report.docx"
try:
    os.remove(out_file)
except:
    pass
doc_generator = DocGenerator(json_file, out_file)
doc_generator.generate()
